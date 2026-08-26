"""RPAForge IDP Library - Intelligent Document Processing parsers."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from rpaforge.core.activity import activity, library, output, param, tags
from rpaforge_libraries.i18n import _
from rpaforge_libraries.IDP.exceptions import (
    IDPDependencyError,
    IDPError,
    IDPParseError,
)

__all__ = [
    "IDPDependencyError",
    "IDPError",
    "IDPParseError",
    "IDP",
]

logger = logging.getLogger("rpaforge.idp")


_INSTALL_HINT = _("Install it with: pip install 'rpaforge-libraries[idp]'")


def _require_pypdf() -> Any:
    """Import pypdf lazily, raising an actionable typed error when missing."""
    try:
        import pypdf
    except ImportError as err:
        raise IDPDependencyError(
            _("pypdf is required to parse PDF files. ") + _INSTALL_HINT
        ) from err
    return pypdf


def _require_docx() -> Any:
    """Import python-docx lazily, raising an actionable typed error."""
    try:
        import docx
    except ImportError as err:
        raise IDPDependencyError(
            _("python-docx is required to parse DOCX files. ") + _INSTALL_HINT
        ) from err
    return docx


def _require_pillow() -> Any:
    """Import pillow lazily, raising an actionable typed error when missing."""
    try:
        from PIL import Image
    except ImportError as err:
        raise IDPDependencyError(
            _("pillow is required to load images. ") + _INSTALL_HINT
        ) from err
    return Image


def _require_document(path: str | Path) -> Path:
    """Validate that *path* points to a non-empty file.

    :param path: Candidate document path.
    :returns: The resolved path.
    :raises FileNotFoundError: If the file does not exist.
    :raises IDPParseError: If the file is empty.
    """
    document_path = Path(path)
    if not document_path.is_file():
        raise FileNotFoundError(_("File not found: {path}", path=str(document_path)))
    if document_path.stat().st_size == 0:
        raise IDPParseError(_("Document '{path}' is empty", path=str(document_path)))
    return document_path


def _parse_page_selection(selection: str | None, total: int) -> list[int]:
    """Resolve a 1-based page selection into sorted 0-based page indices.

    Accepts ``None``/empty (all pages), single pages (``"3"``), ranges
    (``"1-4"``), and comma-separated combinations (``"1,3-5"``).

    :param selection: Raw selection string or ``None``.
    :param total: Total number of pages in the document.
    :returns: Sorted list of 0-based page indices.
    :raises ValueError: If the selection is malformed or out of range.
    """
    if selection is None or not str(selection).strip():
        return list(range(total))
    picked: set[int] = set()
    for part in str(selection).split(","):
        part = part.strip()
        if not part:
            continue
        bounds = part.split("-", 1)
        try:
            if len(bounds) == 2:
                start = int(bounds[0].strip())
                end = int(bounds[1].strip())
                picked.update(range(start, end + 1))
            else:
                picked.add(int(part))
        except ValueError as err:
            raise ValueError(_("Invalid page range: {part}", part=part)) from err
    out_of_range = [page for page in sorted(picked) if page < 1 or page > total]
    if out_of_range:
        raise ValueError(
            _(
                "Page selection {pages} is out of range 1-{total}",
                pages=", ".join(str(page) for page in out_of_range),
                total=total,
            )
        )
    return sorted(page - 1 for page in picked)


def _unique_header(cells: list[str]) -> list[str]:
    """Return table header names de-duplicated with numeric suffixes."""
    seen: dict[str, int] = {}
    unique: list[str] = []
    for name in cells:
        count = seen.get(name, 0)
        unique.append(name if count == 0 else f"{name}_{count + 1}")
        seen[name] = count + 1
    return unique


@library(name="IDP", category="Documents", icon="📄")
class IDP:
    """Intelligent Document Processing library - PDF, DOCX and image parsing.

    All activities return plain pipeline documents built from dicts, lists,
    strings and numbers only, so they cross the stateful-subprocess boundary
    cleanly. Optional parser dependencies are imported lazily; install them
    with ``pip install 'rpaforge-libraries[idp]'``.
    """

    @activity(name="Parse PDF", category="IDP", timeout_ms=120000)
    @tags("idp", "pdf", "document", "parse")
    @output("Pipeline document dict with page-wise text and layout hints")
    @param("path", type="string", description="Path to the PDF file.")
    @param(
        "pages",
        type="string",
        description="Optional 1-based pages to parse, e.g. '1,3-5'. All pages by default.",
    )
    def parse_pdf(self, path: str | Path, pages: str | None = None) -> dict[str, Any]:
        """Parse a PDF's native text layer into a pipeline document dict.

        Each entry of ``pages`` carries the extracted text plus layout hints
        (media-box width and height in points).

        :param path: Path to the PDF file.
        :param pages: Optional 1-based page selection, e.g. ``"1,3-5"``.
        :returns: Document dict with ``source``, ``page_count`` and
            ``pages`` (list of ``{"text": str, "width": int, "height": int}``).
        :raises FileNotFoundError: If the file does not exist.
        :raises IDPParseError: If the file is empty or corrupt.
        :raises ValueError: If the page selection is invalid.
        """
        pypdf = _require_pypdf()
        pdf_path = _require_document(path)
        try:
            reader = pypdf.PdfReader(str(pdf_path))
        except Exception as err:
            raise IDPParseError(
                _(
                    "Failed to open PDF '{path}'. The file may be corrupt: {error}",
                    path=str(pdf_path),
                    error=err,
                )
            ) from err

        indices = _parse_page_selection(pages, len(reader.pages))
        document_pages: list[dict[str, Any]] = []
        for index in indices:
            pdf_page = reader.pages[index]
            try:
                text = pdf_page.extract_text() or ""
                width = int(round(float(pdf_page.mediabox.width or 0)))
                height = int(round(float(pdf_page.mediabox.height or 0)))
            except Exception as err:
                raise IDPParseError(
                    _(
                        "Failed to extract text from page {page} of '{path}': {error}",
                        page=index + 1,
                        path=str(pdf_path),
                        error=err,
                    )
                ) from err
            document_pages.append(
                {"text": text, "width": width, "height": height, "number": index + 1}
            )

        logger.info(
            _(
                "Parsed PDF '{path}': {count} page(s)",
                path=str(pdf_path),
                count=len(document_pages),
            )
        )
        return {
            "source": str(pdf_path),
            "page_count": len(reader.pages),
            "pages": document_pages,
        }

    @activity(name="Parse DOCX", category="IDP", timeout_ms=120000)
    @tags("idp", "docx", "document", "parse")
    @output("Pipeline document dict with paragraphs and tables")
    @param("path", type="string", description="Path to the DOCX file.")
    def parse_docx(self, path: str | Path) -> dict[str, Any]:
        """Parse a DOCX file into a pipeline document dict.

        Returns body paragraphs plus every table as a list of row dicts keyed
        by the first table row (header). The synthetic single-page ``pages``
        entry holds the full paragraph text and section page dimensions in
        points.

        :param path: Path to the DOCX file.
        :returns: Document dict with ``paragraphs``, ``tables`` and ``pages``.
        :raises FileNotFoundError: If the file does not exist.
        :raises IDPParseError: If the file is empty or corrupt.
        """
        docx = _require_docx()
        docx_path = _require_document(path)
        try:
            word_document = docx.Document(str(docx_path))
        except Exception as err:
            raise IDPParseError(
                _(
                    "Failed to open DOCX '{path}'. The file may be empty or corrupt: {error}",
                    path=str(docx_path),
                    error=err,
                )
            ) from err

        paragraphs = [para.text for para in word_document.paragraphs]

        tables: list[list[dict[str, str]]] = []
        for table in word_document.tables:
            rows_out: list[dict[str, str]] = []
            header: list[str] | None = None
            for row_index, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_index == 0:
                    header = _unique_header(cells)
                    continue
                rows_out.append(dict(zip(header or [], cells, strict=False)))
            tables.append(rows_out)

        width = height = 0
        if word_document.sections:
            section = word_document.sections[0]
            if section.page_width is not None:
                width = int(round(section.page_width.pt))
            if section.page_height is not None:
                height = int(round(section.page_height.pt))

        logger.info(
            _(
                "Parsed DOCX '{path}': {paras} paragraph(s), {tbls} table(s)",
                path=str(docx_path),
                paras=len(paragraphs),
                tbls=len(tables),
            )
        )
        return {
            "source": str(docx_path),
            "paragraphs": paragraphs,
            "tables": tables,
            "pages": [
                {
                    "text": "\n".join(paragraphs),
                    "width": width,
                    "height": height,
                    "number": 1,
                }
            ],
        }

    @activity(name="Load Image", category="IDP", timeout_ms=60000)
    @tags("idp", "image", "document", "tiff", "png", "jpeg")
    @output("Pipeline document dict normalized from the image")
    @param(
        "path",
        type="string",
        description="Path to the image file (TIFF, PNG, JPEG, or any Pillow-readable image).",
    )
    def load_image(self, path: str | Path) -> dict[str, Any]:
        """Normalize an image file into a pipeline document dict.

        TIFF/PNG/JPEG (and any other Pillow-readable image) become one page
        per frame with empty native text - OCR runs separately. Multi-frame
        TIFFs produce one page per frame.

        :param path: Path to the image file.
        :returns: Document dict with ``format``, ``mode``, ``frames`` and
            ``pages`` (list of ``{"text": str, "width": int, "height": int}``).
        :raises FileNotFoundError: If the file does not exist.
        :raises IDPParseError: If the file is empty or not a readable image.
        """
        pillow_image = _require_document(path)
        image_module = _require_pillow()
        try:
            with image_module.open(pillow_image) as opened:
                frame_count = max(1, int(getattr(opened, "n_frames", 1)))
                source_format = (opened.format or "").lower() or "image"
                mode = opened.mode
                pages: list[dict[str, Any]] = []
                for frame in range(frame_count):
                    if frame:
                        opened.seek(frame)
                    width, height = opened.size
                    pages.append(
                        {
                            "text": "",
                            "width": int(width),
                            "height": int(height),
                            "number": frame + 1,
                        }
                    )
        except IDPError:
            raise
        except Exception as err:
            raise IDPParseError(
                _(
                    "Failed to load image '{path}'. The file may be corrupt: {error}",
                    path=str(pillow_image),
                    error=err,
                )
            ) from err

        logger.info(
            _(
                "Loaded image '{path}' ({fmt}, {frames} frame(s))",
                fmt=source_format,
                frames=frame_count,
            ),
            path=str(pillow_image),
        )
        return {
            "source": str(pillow_image),
            "format": source_format,
            "mode": mode,
            "frames": frame_count,
            "pages": pages,
        }

    @activity(name="Parse Scanned PDF", category="IDP", timeout_ms=300000)
    @tags("idp", "pdf", "ocr", "scan", "vlm", "hybrid")
    @output(
        "Pipeline document dict with per-page text, engine and confidence provenance"
    )
    @param("path", type="string", description="Path to the scanned PDF file.")
    @param(
        "pages",
        type="string",
        description="Optional 1-based pages to process, e.g. '1,3-5'. All pages by default.",
    )
    @param(
        "min_confidence",
        type="float",
        description="Mean word-confidence (0..1) below which a page escalates to VLM.",
    )
    @param(
        "vlm_fallback",
        type="boolean",
        description="Re-read low-confidence pages via the multimodal LLM provider (sends page images externally).",
    )
    def parse_scanned_pdf(
        self,
        path: str | Path,
        pages: str | None = None,
        min_confidence: float = 0.75,
        vlm_fallback: bool = True,
    ) -> dict[str, Any]:
        """OCR a scanned PDF with local Tesseract plus optional VLM fallback.

        Each page is rasterized and read by Tesseract ``image_to_data``;
        the page-level mean word-confidence gates escalation. Below
        *min_confidence* (and with *vlm_fallback* on) the page image is
        re-read by the configured multimodal LLM provider - page content
        then leaves this machine, which is logged and surfaced in
        ``warnings``.

        :param path: Path to the scanned PDF.
        :param pages: Optional 1-based page selection, e.g. ``"1,3-5"``.
        :param min_confidence: Escalation threshold in 0..1.
        :param vlm_fallback: Enable the external VLM re-read pass.
        :returns: Document dict with ``source``, ``page_count``, ``pages``
            (each ``{"number", "text", "engine": "tesseract"|"vlm",
            "confidence", "width", "height"}``) and ``warnings``.
        :raises FileNotFoundError: If the file does not exist.
        :raises IDPParseError: If the file is empty.
        :raises IDPDependencyError: If pypdfium2/pytesseract/pillow are missing.
        """
        from rpaforge_libraries.IDP.ocr_pipeline import ocr_scanned_document

        pdf_path = _require_document(path)
        reader_page_count = self._pdf_page_count(pdf_path)
        indices = _parse_page_selection(pages, reader_page_count)
        document = ocr_scanned_document(
            pdf_path,
            indices,
            reader_page_count,
            min_confidence=min_confidence,
            vlm_fallback=vlm_fallback,
        )
        logger.info(
            _(
                "Hybrid OCR finished for '{path}': {count} page(s)",
                path=str(pdf_path),
                count=len(document["pages"]),
            )
        )
        return document

    @staticmethod
    def _pdf_page_count(pdf_path: Path) -> int:
        """Return total page count using pypdf."""
        pypdf = _require_pypdf()
        try:
            reader = pypdf.PdfReader(str(pdf_path))
        except Exception as err:
            raise IDPParseError(
                _(
                    "Failed to open PDF '{path}'. The file may be corrupt: {error}",
                    path=str(pdf_path),
                    error=err,
                )
            ) from err
        return len(reader.pages)

    @activity(name="Get Extraction Schema", category="IDP", timeout_ms=30000)
    @tags("idp", "schema", "llm", "extraction", "invoice", "receipt")
    @output("JSON Schema dict for Extract Structured Data")
    @param(
        "doc_type",
        type="string",
        options=["invoice", "receipt", "purchase_order"],
        description="Bundled document schema to load.",
    )
    def get_extraction_schema(self, doc_type: str) -> dict[str, Any]:
        """Load a pre-built extraction schema bundled with the library.

        Schemas ship as package data (importlib.resources) and pair with
        the AI activity ``Extract Structured Data`` for one-activity
        document parsing:

            schema = idp.get_extraction_schema("invoice")
            result = ai.extract_structured_data(text, schema)

        :param doc_type: ``invoice``, ``receipt`` or ``purchase_order``
            (case-insensitive).
        :returns: JSON Schema dict.
        :raises KeyError: For unknown document types.
        """
        from rpaforge_libraries.IDP.extraction_schemas import load_schema

        schema = load_schema(doc_type)
        logger.info(
            _(
                "Loaded extraction schema '{doc_type}' v{version}",
                doc_type=doc_type,
                version=schema.get("version", "?"),
            )
        )
        return schema

    @activity(name="Extract Tables", category="IDP", timeout_ms=120000)
    @tags("idp", "table", "line-items", "ocr", "alignment")
    @output("List of table dicts with headers, rows and per-cell confidence")
    @param("doc", type="dict", description="Pipeline document from IDP parsers or OCR.")
    @param(
        "strategy",
        type="string",
        options=["alignment", "whitespace"],
        description="Column recovery strategy; alignment needs OCR word boxes.",
    )
    def extract_tables(
        self,
        doc: dict[str, Any],
        strategy: str = "alignment",
    ) -> list[dict[str, Any]]:
        """Extract tabular regions from a pipeline document.

        ``alignment`` clusters OCR word x-starts into columns (needs the
        ``words`` entries produced by Parse Scanned PDF); pages without
        coordinates automatically fall back to whitespace splitting.

        :param doc: Pipeline document dict.
        :param strategy: ``"alignment"`` or ``"whitespace"``.
        :returns: List of tables: ``{"page", "strategy", "headers", "rows",
            "confidence" (per-cell matrix or None), "low_confidence_cells"}``.
        """
        from rpaforge_libraries.IDP.tables import extract_tables as _extract

        if not isinstance(doc, dict):
            raise IDPParseError(_("Extract Tables expects a pipeline document dict."))
        tables = _extract(doc, strategy=strategy)
        logger.info(_("Extracted {count} table(s) from document", count=len(tables)))
        return tables

    @activity(name="Table To Records", category="IDP", timeout_ms=60000)
    @tags("idp", "table", "records", "dataframes", "excel")
    @output("List of row dicts keyed by header names")
    @param("table", type="dict", description="Table dict from Extract Tables.")
    @param("headers", type="list", description="Optional column-name override.")
    @param(
        "include_confidence",
        type="boolean",
        description="Add per-cell confidence as '<column>_confidence' keys.",
    )
    def table_to_records(
        self,
        table: dict[str, Any],
        headers: list[str] | None = None,
        include_confidence: bool = False,
    ) -> list[dict[str, Any]]:
        """Convert an extracted table into DataFrame/Excel-ready records.

        Mirrors the web-scraper-to-Excel flow output shape so results feed
        directly into DataFrames/Excel activities.

        :param table: Table dict from Extract Tables.
        :param headers: Optional column-name override.
        :param include_confidence: Attach per-cell confidences when present.
        :returns: List of row dicts keyed by header names.
        """
        from rpaforge_libraries.IDP.tables import table_to_records as _convert

        if not isinstance(table, dict):
            raise IDPParseError(_("Table To Records expects an extracted table dict."))
        records = _convert(
            table,
            headers=headers,
            include_confidence=include_confidence,
        )
        logger.info(_("Converted table to {count} record(s)", count=len(records)))
        return records
