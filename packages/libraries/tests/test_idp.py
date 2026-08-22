"""Tests for IDP library (Intelligent Document Processing)."""

from __future__ import annotations

import builtins
import json
from pathlib import Path

import pytest

from rpaforge_libraries.IDP import IDP, IDPDependencyError, IDPParseError


def make_minimal_pdf(text: str = "Hello IDP", page_count: int = 1) -> bytes:
    """Build a tiny valid PDF with *page_count* pages containing *text*.

    Hand-crafted bytes - no reportlab or other generator dependency.
    """
    objects: list[bytes] = []
    kid_refs = " ".join(f"{3 + i} 0 R" for i in range(page_count))
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(
        f"<< /Type /Pages /Kids [{kid_refs}] /Count {page_count} >>".encode()
    )
    content_ref = 3 + page_count
    for index in range(page_count):
        stream = f"BT /F1 18 Tf 72 720 Td ({text} {index + 1}) Tj ET".encode("latin-1")
        # Layout: objs 1-2 catalog/pages, 3..N pages, N+1..M streams, M+1 font.
        page_obj = (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {content_ref + page_count} 0 R >> >> "
            f"/Contents {content_ref + index} 0 R >>"
        )
        objects.append(page_obj.encode())
        objects.append(
            b"<< /Length "
            + str(len(stream)).encode()
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    )
    out += trailer.encode()
    return bytes(out)


def write_pdf(path: Path, text: str = "Hello IDP", page_count: int = 1) -> Path:
    """Write a minimal generated PDF to *path* and return it."""
    path.write_bytes(make_minimal_pdf(text=text, page_count=page_count))
    return path


def hide_module(monkeypatch: pytest.MonkeyPatch, name: str) -> None:
    """Force importing *name* to raise ImportError even if installed."""
    real_import = builtins.__import__

    def fake_import(module_name, globals=None, locals=None, fromlist=(), level=0):
        if module_name == name or module_name.startswith(name + "."):
            raise ImportError(f"No module named '{module_name}'")
        return real_import(module_name, globals, locals, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", fake_import)


class TestIDPMeta:
    """Tests for library registration metadata."""

    def test_import_library(self):
        lib = IDP()
        assert lib is not None

    def test_library_is_decorated(self):
        assert hasattr(IDP, "_library_meta")
        assert IDP._library_name == "IDP"

    def test_activities_registered(self):
        from rpaforge.core.activity import ACTIVITY_REGISTRY

        for activity_id in ("parse_pdf", "parse_docx", "load_image"):
            meta = ACTIVITY_REGISTRY.get(f"IDP.{activity_id}")
            assert meta is not None, f"missing activity {activity_id}"
            assert meta.library == "IDP"


class TestParsePDF:
    """Tests for the Parse PDF activity (requires the idp extra)."""

    @pytest.fixture(autouse=True)
    def _require_pypdf(self):
        pytest.importorskip("pypdf", reason="pypdf not installed (idp extra required)")

    def test_parses_generated_fixture(self, tmp_path):
        pdf_path = write_pdf(tmp_path / "hello.pdf")
        result = IDP().parse_pdf(pdf_path)

        assert result["page_count"] == 1
        assert len(result["pages"]) == 1
        page = result["pages"][0]
        assert "Hello IDP" in page["text"]
        assert isinstance(page["width"], int)
        assert isinstance(page["height"], int)
        assert (page["width"], page["height"]) == (612, 792)

    def test_document_object_is_json_safe(self, tmp_path):
        pdf_path = write_pdf(tmp_path / "doc.pdf")
        result = IDP().parse_pdf(pdf_path)

        serialized = json.dumps(result)
        assert "Hello IDP" in serialized

    def test_page_selection_subset(self, tmp_path):
        pdf_path = write_pdf(tmp_path / "multi.pdf", page_count=3)
        result = IDP().parse_pdf(pdf_path, pages="1,3")

        assert result["page_count"] == 3
        assert [page["number"] for page in result["pages"]] == [1, 3]
        assert "Hello IDP 1" in result["pages"][0]["text"]
        assert "Hello IDP 3" in result["pages"][1]["text"]

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            IDP().parse_pdf(tmp_path / "missing.pdf")

    def test_empty_file_raises_typed_error(self, tmp_path):
        empty = tmp_path / "empty.pdf"
        empty.write_bytes(b"")

        with pytest.raises(IDPParseError):
            IDP().parse_pdf(empty)

    def test_corrupt_file_raises_typed_error(self, tmp_path):
        corrupt = tmp_path / "corrupt.pdf"
        corrupt.write_bytes(b"%PDF-1.4 this is not a real pdf body")

        with pytest.raises(IDPParseError):
            IDP().parse_pdf(corrupt)


class TestParseDOCX:
    """Tests for the Parse DOCX activity (requires the idp extra)."""

    @pytest.fixture(autouse=True)
    def _require_docx(self):
        pytest.importorskip(
            "docx", reason="python-docx not installed (idp extra required)"
        )

    def _make_docx(self, path: Path) -> Path:
        import docx

        document = docx.Document()
        document.add_paragraph("Invoice header line")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Price"
        table.cell(1, 0).text = "Widget"
        table.cell(1, 1).text = "9.99"
        document.save(str(path))
        return path

    def test_paragraphs_and_tables(self, tmp_path):
        docx_path = self._make_docx(tmp_path / "invoice.docx")
        result = IDP().parse_docx(docx_path)

        assert any("Invoice header line" in p for p in result["paragraphs"])
        assert result["tables"] == [[{"Item": "Widget", "Price": "9.99"}]]
        page = result["pages"][0]
        assert set(page) >= {"text", "width", "height"}
        assert isinstance(page["width"], int)
        assert isinstance(page["height"], int)

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            IDP().parse_docx(tmp_path / "missing.docx")

    def test_empty_file_raises_typed_error(self, tmp_path):
        empty = tmp_path / "empty.docx"
        empty.write_bytes(b"")

        with pytest.raises(IDPParseError):
            IDP().parse_docx(empty)

    def test_corrupt_file_raises_typed_error(self, tmp_path):
        corrupt = tmp_path / "corrupt.docx"
        corrupt.write_bytes(b"PK\x03\x04 definitely not a docx zip")

        with pytest.raises(IDPParseError):
            IDP().parse_docx(corrupt)


class TestLoadImage:
    """Tests for the Load Image activity."""

    def _write_png(self, path: Path, size: tuple[int, int] = (32, 16)) -> Path:
        from PIL import Image

        Image.new("RGB", size, color=(200, 10, 10)).save(path, format="PNG")
        return path

    def test_load_png(self, tmp_path):
        png_path = self._write_png(tmp_path / "scan.png", size=(32, 16))
        result = IDP().load_image(png_path)

        assert result["format"] == "png"
        assert result["frames"] == 1
        assert result["pages"] == [{"text": "", "width": 32, "height": 16, "number": 1}]

    def test_load_multipage_tiff(self, tmp_path):
        from PIL import Image

        tiff_path = tmp_path / "frames.tiff"
        frame_a = Image.new("RGB", (10, 20), color=(255, 0, 0))
        frame_b = Image.new("RGB", (30, 40), color=(0, 255, 0))
        frame_a.save(tiff_path, format="TIFF", save_all=True, append_images=[frame_b])

        result = IDP().load_image(tiff_path)

        assert result["format"] == "tiff"
        assert result["frames"] == 2
        assert [(p["width"], p["height"]) for p in result["pages"]] == [
            (10, 20),
            (30, 40),
        ]

    def test_load_jpeg(self, tmp_path):
        from PIL import Image

        jpeg_path = tmp_path / "photo.jpeg"
        Image.new("RGB", (64, 48)).save(jpeg_path, format="JPEG")

        result = IDP().load_image(jpeg_path)
        assert result["format"] == "jpeg"

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            IDP().load_image(tmp_path / "missing.png")

    def test_empty_file_raises_typed_error(self, tmp_path):
        empty = tmp_path / "empty.png"
        empty.write_bytes(b"")

        with pytest.raises(IDPParseError):
            IDP().load_image(empty)

    def test_not_an_image_raises_typed_error(self, tmp_path):
        fake = tmp_path / "fake.png"
        fake.write_bytes(b"this is plain text, not an image")

        with pytest.raises(IDPParseError):
            IDP().load_image(fake)


class TestPageSelection:
    """Unit tests for 1-based page selection parsing."""

    @pytest.mark.parametrize(
        ("selection", "total", "expected"),
        [
            (None, 3, [0, 1, 2]),
            ("", 2, [0, 1]),
            ("2", 3, [1]),
            ("1,3", 5, [0, 2]),
            ("1-3", 5, [0, 1, 2]),
            ("1, 3-4 ,5", 5, [0, 2, 3, 4]),
        ],
    )
    def test_selections(self, selection, total, expected):
        from rpaforge_libraries.IDP.library import _parse_page_selection

        assert _parse_page_selection(selection, total) == expected

    @pytest.mark.parametrize(("selection", "total"), [("9", 3), ("0", 3), ("1-99", 3)])
    def test_out_of_range_raises(self, selection, total):
        from rpaforge_libraries.IDP.library import _parse_page_selection

        with pytest.raises(ValueError, match="range"):
            _parse_page_selection(selection, total)

    def test_garbage_raises(self):
        from rpaforge_libraries.IDP.library import _parse_page_selection

        with pytest.raises(ValueError, match="Invalid"):
            _parse_page_selection("abc", 3)


class TestDependencyErrors:
    """Lazy imports raise actionable typed errors when extras are missing."""

    def test_dependency_error_is_import_error(self):
        assert issubclass(IDPDependencyError, ImportError)

    def test_parse_pdf_missing_pypdf(self, tmp_path, monkeypatch):
        hide_module(monkeypatch, "pypdf")
        pdf_path = write_pdf(tmp_path / "doc.pdf")

        with pytest.raises(IDPDependencyError) as excinfo:
            IDP().parse_pdf(pdf_path)

        assert "rpaforge-libraries[idp]" in str(excinfo.value)

    def test_parse_docx_missing_docx(self, tmp_path, monkeypatch):
        hide_module(monkeypatch, "docx")
        target = tmp_path / "doc.docx"
        target.write_bytes(b"stub")

        with pytest.raises(IDPDependencyError) as excinfo:
            IDP().parse_docx(target)

        assert "rpaforge-libraries[idp]" in str(excinfo.value)

    def test_load_image_missing_pillow(self, tmp_path, monkeypatch):
        hide_module(monkeypatch, "PIL")
        png_path = tmp_path / "img.png"
        png_path.write_bytes(b"stub")

        with pytest.raises(IDPDependencyError) as excinfo:
            IDP().load_image(png_path)

        assert "rpaforge-libraries[idp]" in str(excinfo.value)
